# -*- coding: utf-8 -*-
# STREAMLIT COM MAPA INTERATIVO FOLIUM - VERSÃO FINAL CORRIGIDA PARA DOCX (2024)

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime, date
import io
import warnings

# === IMPORTAÇÕES NOVAS PARA DOCX ===
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
# ==================================


import folium
from folium import plugins
from streamlit_folium import folium_static
import tempfile
import base64
import math
import requests 


warnings.filterwarnings("ignore")

# Configuração de estilo
plt.style.use("default")
sns.set_palette("husl")

# ==============================================================================
# ⚠️ FUNÇÃO ESSENCIAL: CARREGAMENTO E PRÉ-PROCESSAMENTO DOS DADOS
# ==============================================================================
@st.cache_data
def load_data():
    """
    Simula o carregamento e pré-processamento dos dados de acidentes rodoviários.
    Corrigido para o intervalo 2007-2023.
    """
    st.info("⚠️ Usando dados simulados para demonstração no período 2007-2023.")
   

    np.random.seed(42)
    # Aumentar o número de linhas para simular um dataset maior
    N_ROWS = 50000 

    # Lista de estados brasileiros e suas coordenadas
    estados_coords = {
        "SP": [-23.5489, -46.6388], "MG": [-19.9167, -43.9345],
        "PR": [-25.4284, -49.2733], "RS": [-30.0346, -51.2177],
        "SC": [-27.5954, -48.5480], "RJ": [-22.9068, -43.1729],
        "BA": [-12.9714, -38.5014], "GO": [-16.6809, -49.2533],
        "PE": [-8.0476, -34.8770], "CE": [-3.7172, -38.5434],
        "PA": [-1.4554, -48.4902], "MA": [-2.5387, -44.2823],
        "MT": [-15.6014, -56.0979], "MS": [-20.4428, -54.6464],
        "PI": [-5.0920, -42.8038], "RN": [-5.7945, -35.2110],
        "AL": [-9.6658, -35.7350], "SE": [-10.9472, -37.0731],
        "RO": [-8.7612, -63.9005], "TO": [-10.2491, -48.3243],
        "AC": [-9.9740, -67.8076], "AM": [-3.1190, -60.0217],
        "RR": [2.8235, -60.6758], "AP": [0.0349, -51.0664],
        "DF": [-15.7797, -47.9297],
    }

    # Criar DataFrame
    df = pd.DataFrame(
        {
            "id": range(1, N_ROWS + 1),
            "uf": np.random.choice(list(estados_coords.keys()), N_ROWS),
            "tipo_acidente": np.random.choice(
                [
                    "Colisão Traseira", "Saída de Pista", "Colisão Frontal",
                    "Tombamento", "Atropelamento", "Danos Materiais"
                ],
                N_ROWS,
            ),
            "mortos": np.random.poisson(0.1, N_ROWS),
            "feridos_graves": np.random.poisson(0.3, N_ROWS),
            "feridos_leves": np.random.poisson(1.0, N_ROWS),
            "ilesos": np.random.poisson(2.0, N_ROWS),
            "br": np.random.choice(
                [
                    "BR-101", "BR-116", "BR-040", "BR-381", "BR-153", 
                    "BR-364", "BR-262", "BR-230", "BR-050", "BR-470"
                ],
                N_ROWS,
            ),
            "km": np.random.randint(1, 800, N_ROWS),
        }
    )

    # Corrigido: Adicionar datas no intervalo 2007-2023
    dates = pd.date_range("2007-01-01", "2023-12-31", periods=N_ROWS)
    df["data_inversa"] = np.random.choice(dates, N_ROWS)
    df["ano"] = df["data_inversa"].dt.year
    df["mes"] = df["data_inversa"].dt.month
    df["dia"] = df["data_inversa"].dt.day

    # Horários aleatórios
    horarios = [f"{h:02d}:{m:02d}:00" for h in range(24) for m in range(0, 60, 30)]
    df["horario"] = np.random.choice(horarios, N_ROWS)

    # Dias da semana em português
    dias_map = {0: "Segunda-feira", 1: "Terça-feira", 2: "Quarta-feira", 
                3: "Quinta-feira", 4: "Sexta-feira", 5: "Sábado", 6: "Domingo"}
    df["dia_semana"] = df["data_inversa"].dt.dayofweek.map(dias_map)

    # Adicionar coordenadas com dispersão (jitter)
    def get_coordinates(uf):
        lat, lon = estados_coords[uf]
        return lat + np.random.normal(0, 0.3), lon + np.random.normal(0, 0.3)

    coords = df["uf"].apply(get_coordinates)
    df["latitude"] = coords.apply(lambda x: x[0])
    df["longitude"] = coords.apply(lambda x: x[1])

    return df, estados_coords

# ==============================================================================
# FUNÇÕES AUXILIARES
# ==============================================================================
def parse_coordinate(coord):
    """
    Função robusta para converter coordenadas de vários formatos
    """
    if pd.isna(coord):
        return None

    try:
        # Se for string, fazer limpeza
        if isinstance(coord, str):
            coord = coord.strip().replace(',', '.')

        coord_float = float(coord)

        # Validar intervalos razoáveis para coordenadas brasileiras
        if (-35 <= coord_float <= 5) or (-75 <= coord_float <= -30): 
            return coord_float
        else:
            return None

    except (ValueError, TypeError, IndexError):
        return None

# ==============================================================================
# CLASSES DE ANÁLISE E GERAÇÃO DE RELATÓRIOS - VERSÃO DOCX
# ==============================================================================
class DOCXReportGenerator:
    """Gera o relatório no formato Microsoft Word (.docx)"""
    def __init__(self):
        # Inicializa o documento Word
        self.document = Document()
        # Define o tamanho padrão da imagem em polegadas (5.0") para melhor visualização no Word
        self.default_img_width = Inches(5.0) 

    def add_image(self, image_path, width=None):
        width = width if width is not None else self.default_img_width
        try:
            # Adiciona a imagem e usa a largura em Inches
            self.document.add_picture(image_path, width=width)
            self.document.add_paragraph().add_run().add_break() # Adiciona uma quebra de linha/espaçador
        except Exception as e:
            self.add_paragraph(f"Erro ao carregar imagem: {str(e)}")

    def add_heading(self, text, level=1):
        # Mapeia nível 1 para Heading 1 (Word) e nível 2 para Heading 2 (Word)
        docx_level = 1 if level == 1 else level 
        self.document.add_heading(text, level=docx_level)

    def add_paragraph(self, text, style='Normal'):
        # Trata o negrito do HTML (<b>) para o formato DOCX (Runs)
        p = self.document.add_paragraph(style=style)
        
        if "<b>" in text and "</b>" in text:
            parts = text.split("<b>", 1)
            p.add_run(parts[0]) 
            
            if len(parts) > 1:
                bold_parts = parts[1].split("</b>", 1)
                p.add_run(bold_parts[0]).bold = True 
                if len(bold_parts) > 1:
                    p.add_run(bold_parts[1]) 
        else:
            p.add_run(text)
            
    def add_table(self, data):
        # Data é uma lista de listas: [[header1, header2], [row1_col1, row1_col2], ...]
        if not data:
            return

        table = self.document.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Light Shading' 

        # Adicionar cabeçalhos
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(data[0]):
            hdr_cells[i].text = str(header)
            # Aplicar negrito e centralizar no cabeçalho
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adicionar dados
        for i in range(1, len(data)):
            row_cells = table.rows[i].cells
            for j, cell_data in enumerate(data[i]):
                row_cells[j].text = str(cell_data)
                row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph() # Espaçador

    def generate_docx(self):
        """Salva o documento DOCX em um buffer de memória."""
        buffer = io.BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer

    def build_report(self, analyzer, selections, author, resumo_executivo, insights, figuras_graficos, figuras_mapas):
        """Constrói o relatório DOCX"""
        
        # Cabeçalho
        self.add_heading("RELATÓRIO EXECUTIVO - ANÁLISE DE ACIDENTES RODOVIÁRIOS", 1)
        self.add_paragraph(f"<b>Autor:</b> {author}")
        self.add_paragraph(f"<b>Data de geração:</b> {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        self.add_paragraph(f"<b>Período analisado:</b> {analyzer.df['ano'].min()} - {analyzer.df['ano'].max()}")
        
        # Resumo Executivo
        self.add_heading("RESUMO EXECUTIVO", 2)
        self.add_paragraph(resumo_executivo)
        
        # Métricas Principais
        if selections.get("include_metrics"):
            self.add_heading("PRINCIPAIS MÉTRICAS", 2)
            tabela_metricas = analyzer.create_metrics_table()
            data_metricas = [list(tabela_metricas.columns)] + tabela_metricas.values.tolist()
            self.add_table(data_metricas)
        
        # Análises Gráficas (Salvamento de figuras a partir de buffers)
        for key, fig in figuras_graficos.items():
            if key == "evolution" and selections.get("include_evolution"):
                self.add_heading("EVOLUÇÃO TEMPORAL", 2)
                self.add_paragraph("A análise temporal mostra a evolução dos acidentes ao longo dos anos, permitindo identificar tendências e sazonalidades.")
            elif key == "states" and selections.get("include_states"):
                self.add_heading("ANÁLISE POR ESTADO", 2)
                self.add_paragraph("Comparativo entre estados brasileiros considerando volume de acidentes e taxas de mortalidade.")
            elif key == "types" and selections.get("include_types"):
                self.add_heading("TIPOS DE ACIDENTE", 2)
                self.add_paragraph("Distribuição percentual dos diferentes tipos de acidentes ocorridos no período analisado.")
            elif key == "weekday" and selections.get("include_weekday"):
                self.add_heading("PADRÃO SEMANAL DE ACIDENTES", 2)
                self.add_paragraph("Distribuição dos acidentes por dia da semana, útil para planejamento logístico e operacional.")
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
                # O parâmetro bbox_inches='tight' garante o corte correto
                fig.savefig(tmpfile.name, dpi=300, bbox_inches='tight')
                # A imagem será adicionada com a largura padronizada (5.0*Inches)
                self.add_image(tmpfile.name)
            
        # Análise de Rodovias
        if selections.get("include_highways"):
            self.add_heading("RANKING DE RODOVIAS", 2)
            self.add_paragraph("Identificação das rodovias com maior incidência de acidentes e maiores taxas de mortalidade.")
            tabela_rodovias = analyzer.create_highways_table()
            data_rodovias = [list(tabela_rodovias.columns)] + tabela_rodovias.values.tolist()
            self.add_table(data_rodovias)
        
        # Mapas (Nota: A captura de Folium em PDF é complexa e requer ferramentas externas)
        if selections.get("include_map") or selections.get("include_complete_map"):
            self.add_heading("MAPA DE DISTRIBUIÇÃO GEOGRÁFICA (Visualização na Plataforma)", 2)
            self.add_paragraph("Os mapas interativos não podem ser embutidos em documentos Word ou PDF de forma nativa e interativa. Consulte a plataforma para a visualização completa e dinâmica.")
        
        # Insights e Recomendações
        self.add_heading("INSIGHTS E RECOMENDAÇÕES ESTRATÉGICAS", 2)
        self.add_paragraph(insights)
        
        # Conclusão
        self.add_heading("CONCLUSÃO", 2)
        self.add_paragraph("Este relatório fornece uma visão abrangente da situação dos acidentes rodoviários, identificando padrões, áreas críticas e oportunidades de intervenção para redução de acidentes e melhoria da segurança viária.")


class DataAnalyzer:
    def __init__(self, df, estados_coords):
        # O DF aqui é o df_filtrado da última iteração do main()
        self.df = df
        self.estados_coords = estados_coords
        self.font_settings = {
            'title_size': 14,
            'label_size': 10,
            'legend_size': 9,
            'ticks_size': 8
        }

    def update_font_settings(self, title_size=14, label_size=10, legend_size=9, ticks_size=8):
        """Atualiza configurações de fonte para todos os gráficos"""
        self.font_settings = {
            'title_size': title_size,
            'label_size': label_size,
            'legend_size': legend_size,
            'ticks_size': ticks_size
        }

    # ==================== GRÁFICOS - SEM CACHE ====================
    def create_evolution_chart(_self):
        """Gráfico de evolução temporal dos acidentes rodoviários"""
        df_copy = _self.df.copy()
        anual = (
            df_copy.groupby("ano")
            .agg({"id": "count", "mortos": "sum", "feridos_graves": "sum"})
            .reset_index()
        )

        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(10, 8))

        # Gráfico superior: Evolução do número de acidentes
        ax1.plot(anual["ano"], anual["id"], marker="o", linewidth=2, color="#1f77b4")
        ax1.set_title("Evolução Anual de Acidentes", 
                     fontweight="bold", 
                     fontsize=_self.font_settings['title_size'])
        ax1.set_ylabel("Número de Acidentes", fontsize=_self.font_settings['label_size'])
        ax1.tick_params(axis='both', which='major', labelsize=_self.font_settings['ticks_size'])
        ax1.grid(True, alpha=0.3)

        # Gráfico inferior: Mortos e feridos graves
        ax2.bar(
            anual["ano"] - 0.2,
            anual["mortos"],
            width=0.4,
            label="Mortos",
            alpha=0.7,
            color="#e74c3c",
        )
        ax2.bar(
            anual["ano"] + 0.2,
            anual["feridos_graves"],
            width=0.4,
            label="Feridos Graves",
            alpha=0.7,
            color="#f39c12",
        )
        ax2.legend(fontsize=_self.font_settings['legend_size'])
        ax2.set_xlabel("Ano", fontsize=_self.font_settings['label_size'])
        ax2.set_ylabel("Número de Vítimas", fontsize=_self.font_settings['label_size'])
        ax2.tick_params(axis='both', which='major', labelsize=_self.font_settings['ticks_size'])
        ax2.grid(True, alpha=0.3)

        plt.tight_layout()
        return fig

    def create_states_chart(_self):
        """Análise comparativa entre estados brasileiros"""
        df_copy = _self.df.copy()
        estados = (
            df_copy.groupby("uf").agg({"id": "count", "mortos": "sum"}).reset_index()
        )
        estados["taxa_mortalidade"] = (estados["mortos"] / estados["id"].replace(0, np.nan)) * 100

        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))

        # Gráfico 1: Top 10 estados por número de acidentes
        top10 = estados.nlargest(10, "id")
        ax1.barh(top10["uf"], top10["id"], color="skyblue", alpha=0.8)
        ax1.set_title("Top 10 Estados - Número de Acidentes", 
                     fontweight="bold", 
                     fontsize=_self.font_settings['title_size'])
        ax1.set_xlabel("Quantidade de Acidentes", fontsize=_self.font_settings['label_size'])
        ax1.tick_params(axis='both', which='major', labelsize=_self.font_settings['ticks_size'])

        for i, v in enumerate(top10["id"]):
            ax1.text(v + 3, i, str(v), va="center", fontsize=_self.font_settings['ticks_size'])

        # Gráfico 2: Top 10 estados por taxa de mortalidade
        top_mortalidade = estados.nlargest(10, "taxa_mortalidade").dropna(subset=['taxa_mortalidade'])
        ax2.barh(
            top_mortalidade["uf"],
            top_mortalidade["taxa_mortalidade"],
            color="lightcoral",
            alpha=0.8,
        )
        ax2.set_title("Top 10 Estados - Taxa de Mortalidade", 
                     fontweight="bold", 
                     fontsize=_self.font_settings['title_size'])
        ax2.set_xlabel("Taxa de Mortalidade (%)", fontsize=_self.font_settings['label_size'])
        ax2.tick_params(axis='both', which='major', labelsize=_self.font_settings['ticks_size'])

        for i, v in enumerate(top_mortalidade["taxa_mortalidade"]):
            ax2.text(v + 0.05, i, f"{v:.2f}%", va="center", fontsize=_self.font_settings['ticks_size'])

        plt.tight_layout()
        return fig

    def create_accident_types_chart(_self):
        """Distribuição dos tipos de acidentes ocorridos"""
        df_copy = _self.df.copy()
        tipos = df_copy["tipo_acidente"].value_counts()

        fig, ax = plt.subplots(figsize=(10, 6))
        colors = plt.cm.Set3(np.linspace(0, 1, len(tipos)))
        explode = [0] * len(tipos)
        if not tipos.empty:
            explode[0] = 0.1

        wedges, texts, autotexts = ax.pie(
            tipos.values,
            labels=tipos.index,
            autopct="%1.1f%%",
            startangle=90,
            colors=colors,
            textprops={"fontsize": _self.font_settings['ticks_size']},
            explode=explode,
            shadow=True,
        )

        for autotext in autotexts:
            autotext.set_color("black")
            autotext.set_fontweight("bold")

        ax.set_title(
            "Distribuição por Tipo de Acidente", 
            fontweight="bold", 
            fontsize=_self.font_settings['title_size']
        )
        return fig

    def create_weekday_chart(_self):
        """Padrão de acidentes por dia da semana (Acidentes vs Mortos)"""
        df_copy = _self.df.copy()
        
        # Ordenar dias da semana
        dias_ordem = ["Segunda-feira", "Terça-feira", "Quarta-feira", "Quinta-feira", "Sexta-feira", "Sábado", "Domingo"]
        
        df_copy['dia_semana'] = pd.Categorical(df_copy['dia_semana'], categories=dias_ordem, ordered=True)
        
        semanal = (
            df_copy.groupby("dia_semana")
            .agg({"id": "count", "mortos": "sum"})
            .reset_index()
            .sort_values("dia_semana")
        )
        
        if semanal.empty:
            fig, ax = plt.subplots(figsize=(10, 6))
            ax.set_title("Padrão Semanal de Acidentes (Sem dados)", fontsize=_self.font_settings['title_size'])
            return fig

        fig, ax = plt.subplots(figsize=(10, 6))
        
        # Gráfico de barras para Acidentes
        ax.bar(semanal["dia_semana"], semanal["id"], color="#2ecc71", alpha=0.8)
        
        # Adiciona a linha de Mortos no eixo Y secundário
        ax2 = ax.twinx()
        ax2.plot(semanal["dia_semana"], semanal["mortos"], marker="o", color="#c0392b", linewidth=3)
        
        ax.set_title("Padrão Semanal de Acidentes (Acidentes vs. Mortos)", 
                     fontweight="bold", 
                     fontsize=_self.font_settings['title_size'])
        ax.set_xlabel("Dia da Semana", fontsize=_self.font_settings['label_size'])
        ax.set_ylabel("Número de Acidentes", color="#2ecc71", fontsize=_self.font_settings['label_size'])
        ax2.set_ylabel("Número de Mortos", color="#c0392b", fontsize=_self.font_settings['label_size'])
        
        # Ajustar ticks
        ax.tick_params(axis='x', rotation=45, labelsize=_self.font_settings['ticks_size'])
        ax.tick_params(axis='y', colors='#2ecc71', labelsize=_self.font_settings['ticks_size'])
        ax2.tick_params(axis='y', colors='#c0392b', labelsize=_self.font_settings['ticks_size'])
        
        ax.grid(axis='y', alpha=0.3)
        
        plt.tight_layout()
        return fig

    # ==================== MAPAS INTERATIVOS - CORRIGIDOS ====================
    def create_interactive_map(_self, df_filtrado, sample_size=1000, map_height=500):
        """Mapa interativo com heatmap e marcadores por estado"""
        
        estados_data = (
            df_filtrado.groupby("uf")
            .agg({"id": "count", "mortos": "sum", "feridos_graves": "sum"})
            .reset_index()
        )
        estados_data["taxa_mortalidade"] = (
            estados_data["mortos"] / estados_data["id"]
        ) * 100

        m = folium.Map(
            location=[-15.77972, -47.92972], 
            zoom_start=4,
            tiles='CartoDB positron',
            control_scale=True,
            prefer_canvas=True
        )

        # Adicionar múltiplos tiles
        folium.TileLayer(
            'OpenStreetMap', 
            name='OpenStreetMap', 
            attr='&copy; <a href="https://www.openstreetmap.org/copyright">OpenStreetMap</a> contributors'
        ).add_to(m)
        
        folium.TileLayer(
            'CartoDB dark_matter', 
            name='CartoDB Dark', 
            attr='Tiles by CartoDB, under CC BY 3.0. Data by OpenStreetMap, under ODbL.'
        ).add_to(m)
        
        folium.TileLayer(
            'Stamen Terrain', 
            name='Stamen Terrain', 
            attr='Map tiles by <a href="http://stamen.com">Stamen Design</a>, under <a href="http://creativecommons.org/licenses/by/3.0">CC BY 3.0</a>.'
        ).add_to(m)

        marker_cluster = plugins.MarkerCluster(name="Estados").add_to(m)

        for _, estado in estados_data.iterrows():
            uf = estado["uf"]
            if uf in _self.estados_coords:
                lat, lon = _self.estados_coords[uf]

                taxa = estado["taxa_mortalidade"] if pd.notna(estado["taxa_mortalidade"]) else 0
                if taxa > 2:
                    cor = "red"
                    icon_type = "exclamation-triangle"
                elif taxa > 1:
                    cor = "orange"
                    icon_type = "warning"
                else:
                    cor = "green"
                    icon_type = "info-sign"

                popup_text = f"""
                <div style="font-family: Arial; width: 250px;">
                    <h4 style="color: {cor}; margin-bottom: 10px;">{uf}</h4>
                    <p><b>Acidentes:</b> {estado['id']:,}</p>
                    <p><b>Mortos:</b> {estado['mortos']:,}</p>
                    <p><b>Taxa de Mortalidade:</b> {taxa:.1f}%</p>
                    <p><b>Feridos Graves:</b> {estado['feridos_graves']:,}</p>
                </div>
                """

                folium.Marker(
                    [lat, lon],
                    popup=folium.Popup(popup_text, max_width=300),
                    tooltip=f"🚨 {uf}: {estado['id']} acidentes | {taxa:.1f}% mortalidade",
                    icon=folium.Icon(color=cor, icon=icon_type, prefix='fa'),
                ).add_to(marker_cluster)


        # Heatmap com amostra da área filtrada
        locais = df_filtrado[["latitude", "longitude"]].dropna()
        if len(locais) > 0:
            # Garante que a amostra é feita do df_filtrado atual
            heat_data = [
                [row["latitude"], row["longitude"]]
                for _, row in locais.sample(min(sample_size, len(locais))).iterrows()
            ]
            plugins.HeatMap(
                heat_data, 
                radius=15, 
                blur=10, 
                max_zoom=8,
                gradient={0.4: 'blue', 0.65: 'lime', 1: 'red'}
            ).add_to(m)

        plugins.Fullscreen(position="topright").add_to(m)
        plugins.MiniMap(tile_layer="CartoDB positron", position="bottomright").add_to(m)
        plugins.LocateControl(position="topright").add_to(m)
        plugins.MeasureControl(position="topleft").add_to(m)
        
        folium.LayerControl(collapsed=False).add_to(m)

        return m

    # FUNÇÃO DE MAPA COMPLETO ATUALIZADA COM O CÓDIGO DO MAPA.PY E MARCADORES POSITIVOS
    def create_complete_logistics_map(_self, df_filtrado, sample_size=1000, map_height=600): # Adicionado sample_size
        """Mapa completo com todas as funcionalidades de logística e análise, baseado no Mapa.py avançado"""

        # Mapeamento de variáveis
        df_enriched = df_filtrado.copy()
        coordenadas_estados = _self.estados_coords

        # 1. Calcular dados por estado (estados_acidentes)
        estados_acidentes = (
            df_enriched.groupby("uf")
            .agg({"id": "count", "mortos": "sum", "feridos_graves": "sum"})
            .reset_index()
        )
        estados_acidentes["taxa_mortalidade"] = (
            estados_acidentes["mortos"] / estados_acidentes["id"]
        ) * 100
        estados_acidentes['taxa_mortalidade'] = estados_acidentes['taxa_mortalidade'].fillna(0) # Tratar NaN

        m2 = folium.Map(
            location=[-15.77972, -47.92972],
            zoom_start=4,
            tiles="Esri_WorldImagery", # Tile padrão
            control_scale=True,
            prefer_canvas=True
        )

        # 2. Adicionar múltiplos temas com Atribuição (Attr)
        tiles_config = {
            "Satélite (Real)": {
                "url": "Esri_WorldImagery",
                "attr": "Tiles &copy; Esri &mdash; Source: Esri, i-cubed, USDA, USGS, AEX, GeoEye, Getmapping, Aerogrid, IGN, IGP, UPR-EGP, and the GIS User Community"
            },
            "Modo Claro (CartoDB)": {
                "url": "CartoDB positron",
                "attr": 'Tiles by CartoDB, under CC BY 3.0. Data by OpenStreetMap, under ODbL.'
            },
            "Satélite (OpenStreetMap)": {
                "url": "OpenStreetMap",
                "attr": '&copy; <a href="https://www.openstreetmap.org/copyright">OpenStreetMap</a> contributors'
            },
            "Terreno (Stamen)": {
                "url": "Stamen Terrain",
                "attr": 'Map tiles by <a href="http://stamen.com">Stamen Design</a>, under <a href="http://creativecommons.org/licenses/by/3.0">CC BY 3.0</a>.'
            },
            "Modo Escuro (CartoDB)": {
                "url": "CartoDB dark_matter",
                "attr": 'Tiles by CartoDB, under CC BY 3.0. Data by OpenStreetMap, under ODbL.'
            }
        }

        for name, config in tiles_config.items():
            folium.TileLayer(
                config['url'], 
                name=name, 
                attr=config['attr']
            ).add_to(m2)
            
        # <<< NOVO: INICIALIZAÇÃO DOS FEATURE GROUPS E CLUSTER >>>
        # Grupo para o Mapa de Calor
        fg_heatmap = folium.FeatureGroup(name='🔥 Mapa de Calor (Densidade)', show=False)
        # Grupo para os marcadores de estado agrupados (Cluster)
        marker_cluster_estados = plugins.MarkerCluster(name="📍 Estados (Agrupados)").add_to(m2)


        # 3. Choropleth Map 
        geojson_url = "https://raw.githubusercontent.com/codeforamerica/click_that_hood/master/public/data/brazil-states.geojson"
        geojson_data = None
        try:
            response = requests.get(geojson_url)
            geojson_data = response.json()
        except Exception:
            # Tenta uma alternativa caso a principal falhe
            alternative_url = "https://raw.githubusercontent.com/tbrugz/geodata-br/master/geojson/geojs-100-mun.json"
            try:
                response = requests.get(alternative_url)
                geojson_data = response.json()
            except Exception:
                pass

        if geojson_data is not None:
            folium.Choropleth(
                geo_data=geojson_data,
                name='Taxa de Mortalidade por Estado (Choropleth)',
                data=estados_acidentes,
                columns=['uf', 'taxa_mortalidade'],
                key_on='feature.properties.sigla',
                fill_color='RdPu',
                fill_opacity=0.7,
                line_opacity=0.2,
                legend_name='Taxa de Mortalidade (%) - 2007-2023',
                bins=6,
                nan_fill_color='lightgray',
                show=False # Inicia desativado para não poluir
            ).add_to(m2)
            
        # <<< NOVO: LÓGICA DO MAPA DE CALOR (HEATMAP) DO MAPA BÁSICO >>>
        locais = df_filtrado[["latitude", "longitude"]].dropna()
        if len(locais) > 0:
            heat_data = [
                [row["latitude"], row["longitude"]]
                for _, row in locais.sample(min(sample_size, len(locais))).iterrows()
            ]
            plugins.HeatMap(
                heat_data, 
                radius=15, 
                blur=10, 
                max_zoom=8,
                gradient={0.4: 'blue', 0.65: 'lime', 1: 'red'}
            ).add_to(fg_heatmap)
        fg_heatmap.add_to(m2)


        # 4. Marcadores de Mortalidade por Estado (CircleMarker e agora também MarkerCluster)
        ranking_mortalidade = estados_acidentes['taxa_mortalidade'].rank(method='dense', ascending=False)
        
        # 5. Adicionar Feature Group para Rotas Seguras (Novo ponto 2)
        fg_rotas_seguras = folium.FeatureGroup(name='✅ Rotas Seguras (Mortalidade Zero)', show=False)

        for _, estado in estados_acidentes.iterrows():
            uf = estado["uf"]
            if uf in coordenadas_estados:
                lat, lon = coordenadas_estados[uf]
                taxa = estado["taxa_mortalidade"] if pd.notna(estado["taxa_mortalidade"]) else 0
                posicao = int(ranking_mortalidade[estados_acidentes['uf'] == uf].iloc[0])

                # <<< NOVO: LÓGICA DO MARKERCLUSTER (MAPA BÁSICO) SENDO ADICIONADA AQUI >>>
                if taxa > 2:
                    cor_cluster = "red"
                    icon_type_cluster = "exclamation-triangle"
                elif taxa > 1:
                    cor_cluster = "orange"
                    icon_type_cluster = "warning"
                else:
                    cor_cluster = "green"
                    icon_type_cluster = "info-sign"

                popup_text_cluster = f"""
                <div style="font-family: Arial; width: 250px;">
                    <h4 style="color: {cor_cluster}; margin-bottom: 10px;">{uf}</h4>
                    <p><b>Acidentes:</b> {estado['id']:,}</p>
                    <p><b>Mortos:</b> {estado['mortos']:,}</p>
                    <p><b>Taxa de Mortalidade:</b> {taxa:.1f}%</p>
                    <p><b>Feridos Graves:</b> {estado['feridos_graves']:,}</p>
                </div>
                """
                folium.Marker(
                    [lat, lon],
                    popup=folium.Popup(popup_text_cluster, max_width=300),
                    tooltip=f"📍 {uf}: {estado['id']} acidentes | {taxa:.1f}% mortalidade",
                    icon=folium.Icon(color=cor_cluster, icon=icon_type_cluster, prefix='fa'),
                ).add_to(marker_cluster_estados)
                # <<< FIM DA LÓGICA DO MARKERCLUSTER >>>


                if taxa > 3:
                    icon_color, risco, recomendacao_logistica, cor_recomendacao = "darkred", "MUITO ALTO", "🚨 EVITAR - Alto risco para operações logísticas", "#ff6b6b"
                elif taxa > 2:
                    icon_color, risco, recomendacao_logistica, cor_recomendacao = "red", "ALTO", "⚠️ CUIDADO - Redobrar atenção nas operações", "#ffa726"
                elif taxa > 1:
                    icon_color, risco, recomendacao_logistica, cor_recomendacao = "orange", "MÉDIO", "📋 ATENÇÃO - Implementar protocolos de segurança", "#ffd93d"
                # Marcar Estados com zero mortalidade na camada de Rotas Seguras
                elif taxa == 0:
                    icon_color_seguro, risco_seguro, recomendacao_segura, cor_recomendacao_segura = "green", "BAIXO", "✅ ADEQUADO - Rotas Seguras (Mortalidade Zero)", "#6bcf7f"
                    
                    # Adicionar como marcador na camada de Rotas Seguras
                    popup_seguro = f"""
                    <div style="font-family: Arial; min-width: 250px; background: white; color: black; padding: 15px; border-radius: 8px; border: 2px solid green;">
                        <h4 style="margin: 0 0 10px 0; color: #2ecc71;">✅ {uf} - ROTA SEGURA (ESTADO)</h4>
                        <div style="background: #2ecc71; color: white; padding: 5px; border-radius: 4px; text-align: center; margin-bottom: 10px;">
                            <strong>Taxa de Mortalidade: 0.00%</strong>
                        </div>
                        <table style="width: 100%; font-size: 12px;">
                            <tr><td>📊 Acidentes:</td><td style="text-align: right;"><strong>{estado['id']:,}</strong></td></tr>
                            <tr><td>📦 Recomendação Logística:</td><td style="text-align: right; color: #2ecc71;"><strong>{recomendacao_segura.split(' - ')[0]}</strong></td></tr>
                        </table>
                    </div>
                    """
                    folium.Marker(
                        [lat, lon],
                        popup=folium.Popup(popup_seguro, max_width=300),
                        tooltip=f"✅ {uf}: Rota Segura (0 mortes) | {estado['id']:,} acidentes",
                        icon=folium.Icon(color='green', icon='fa-thumbs-up', prefix='fa'),
                    ).add_to(fg_rotas_seguras)
                    
                    icon_color, risco, recomendacao_logistica, cor_recomendacao = "lightgreen", "BAIXO", "✅ ADEQUADO - Condições aceitáveis para logística", "#6bcf7f"
                else: # Taxa entre 0 e 1% (Baixo Risco)
                    icon_color, risco, recomendacao_logistica, cor_recomendacao = "lightgreen", "BAIXO", "✅ ADEQUADO - Condições aceitáveis para logística", "#6bcf7f"
                
                # Contexto Logístico (Baseado no Mapa.py)
                if uf in ['SP', 'RJ', 'MG', 'ES']:
                    contexto_logistica = "Região com alta densidade logística - múltiplas rotas alternativas disponíveis"
                elif uf in ['PR', 'SC', 'RS']:
                    contexto_logistica = "Infraestrutura rodoviária de boa qualidade - atenção ao inverno"
                elif uf in ['GO', 'MT', 'MS', 'DF']:
                    contexto_logistica = "Grandes distâncias - planejar pontos de apoio e combustível"
                elif uf in ['BA', 'PE', 'CE', 'MA']:
                    contexto_logistica = "Condições climáticas variáveis - verificar previsões"
                else:
                    contexto_logistica = "Consultar condições específicas da região antes do planejamento"

                popup_text = f"""
                <div style="font-family: Arial; min-width: 320px; background: white; color: black; padding: 15px; border-radius: 8px; border: 2px solid {icon_color};">
                    <h4 style="margin: 0 0 10px 0; color: #e74c3c;">{uf} - NÍVEL {risco}</h4>
                    <div style="background: {icon_color}; color: white; padding: 5px; border-radius: 4px; text-align: center; margin-bottom: 10px;">
                        <strong>Taxa de Mortalidade: {taxa:.2f}%</strong>
                    </div>

                    <div style="background: {cor_recomendacao}; color: white; padding: 8px; border-radius: 4px; margin-bottom: 10px; text-align: center;">
                        <strong>📦 RECOMENDAÇÃO LOGÍSTICA</strong>
                    </div>
                    <div style="font-size: 12px; margin-bottom: 10px; padding: 8px; background: #f8f9fa; border-radius: 4px;">
                        {recomendacao_logistica}
                    </div>

                    <table style="width: 100%; font-size: 12px;">
                        <tr><td>📊 Acidentes:</td><td style="text-align: right;"><strong>{estado['id']:,}</strong></td></tr>
                        <tr><td>💀 Mortes:</td><td style="text-align: right;"><strong>{estado['mortos']:,}</strong></td></tr>
                        <tr><td>🏥 Feridos Graves:</td><td style="text-align: right;"><strong>{estado['feridos_graves']:,}</strong></td></tr>
                        <tr><td>🏆 Ranking Mortalidade:</td><td style="text-align: right;"><strong>#{posicao}</strong></td></tr>
                    </table>

                    <div style="margin-top: 10px; padding: 8px; background: #e3f2fd; border-radius: 4px;">
                        <strong>🗺️ CONTEXTO LOGÍSTICO:</strong><br>
                        <span style="font-size: 11px;">{contexto_logistica}</span>
                    </div>
                </div>
                """

                # Crie um FeatureGroup para os Círculos, para poder ligar/desligar
                fg_circulos_risco = folium.FeatureGroup(name='⭕ Estados (Círculos de Risco)', show=True).add_to(m2)
                folium.CircleMarker(
                    location=[lat, lon],
                    radius=10 + (taxa * 3),
                    popup=folium.Popup(popup_text, max_width=400),
                    tooltip=f"⭕ {uf}: {taxa:.2f}% mortalidade | Nível {risco} | 📦 {recomendacao_logistica.split(' - ')[0]}",
                    color=icon_color,
                    fillColor=icon_color,
                    fillOpacity=0.7,
                    weight=2
                ).add_to(fg_circulos_risco)
        
        fg_rotas_seguras.add_to(m2)


        # ... (O restante do código para Acidentes Graves e Todas as Rodovias permanece o mesmo) ...
        # 6. Processar Coordenadas e Adicionar Acidentes Graves (Marker)
        # O parse_coordinate é uma função global, acessível aqui.
        df_coords = df_enriched.copy()
        df_coords['lat_clean'] = df_coords['latitude'].apply(parse_coordinate)
        df_coords['lon_clean'] = df_coords['longitude'].apply(parse_coordinate)
        
        coords_validas = df_coords[
            (df_coords['lat_clean'].notna()) &
            (df_coords['lon_clean'].notna())
        ]

        coluna_br = 'br'
        coluna_km = 'km'

        if coluna_br in df_enriched.columns and coluna_km in df_enriched.columns:
            df_enriched[coluna_br] = df_enriched[coluna_br].fillna('Não informada').astype(str)
            df_enriched[coluna_km] = df_enriched[coluna_km].fillna('Não informado').astype(str)

            acidentes_graves = coords_validas[
                (coords_validas['mortos'] > 0) |
                (coords_validas['feridos_graves'] > 0)
            ].copy()

            amostra_size = min(200, len(acidentes_graves))
            amostra_acidentes = acidentes_graves.head(amostra_size)

            fg_acidentes_graves = folium.FeatureGroup(name='🚨 Acidentes Graves (BR/KM)', show=False)

            for _, acidente in amostra_acidentes.iterrows():
                try:
                    lat = acidente['lat_clean']
                    lon = acidente['lon_clean']

                    if pd.isna(lat) or pd.isna(lon): continue

                    # Limpar valores de BR e KM
                    br = str(acidente[coluna_br]).split('.')[0] if '.' in str(acidente[coluna_br]) else str(acidente[coluna_br])
                    km = str(acidente[coluna_km]).split('.')[0] if '.' in str(acidente[coluna_km]) else str(acidente[coluna_km])

                    mortos = int(acidente['mortos']) if pd.notna(acidente['mortos']) else 0
                    feridos_graves = int(acidente['feridos_graves']) if pd.notna(acidente['feridos_graves']) else 0

                    if mortos > 0:
                        cor, gravidade, impacto_logistica = 'black', "FATAL", "🔴 PARALISAÇÃO TOTAL - Rota deve ser evitada"
                    elif feridos_graves > 0:
                        cor, gravidade, impacto_logistica = 'red', "GRAVE", "🟡 ALTO IMPACTO - Considerar rotas alternativas"
                    else:
                        continue

                    data_acidente = str(acidente['data_inversa']).split(' ')[0] if 'data_inversa' in acidente and pd.notna(acidente['data_inversa']) else "Data não disponível"
                    
                    br_principal = any(main_br in br for main_br in ['101', '116', '381', '040', '153', '364', '262'])
                    contexto_logistica_acidente = "🚚 ROTA PRINCIPAL - Impacto significativo na logística regional" if br_principal else "🛣️ ROTA SECUNDÁRIA - Impacto localizado"
                    
                    horario = acidente.get('horario', 'Não informado')
                    periodo = "⏰ Horário não informado"
                    if horario != 'Não informado':
                        try:
                            hora = int(horario.split(':')[0])
                            periodo = "⏰ HORÁRIO DE PICO (manhã)" if 6 <= hora <= 9 else "⏰ HORÁRIO DE PICO (tarde)" if 16 <= hora <= 19 else "⏰ FORA DE PICO"
                        except:
                            pass

                    popup_html = f"""
                    <div style="font-family: Arial; min-width: 350px; background: white; color: black; padding: 15px; border-radius: 8px; border: 2px solid {cor};">
                        <h4 style="margin: 0 0 10px 0; color: {cor};">🚨 ACIDENTE {gravidade}</h4>
                        <div style="background: {cor}; color: white; padding: 8px; border-radius: 4px; text-align: center; margin-bottom: 10px;">
                            <strong>BR {br} • KM {km}</strong>
                        </div>

                        <div style="background: #ffeb3b; color: #333; padding: 8px; border-radius: 4px; margin-bottom: 10px; text-align: center;">
                            <strong>📦 IMPACTO LOGÍSTICO</strong><br>
                            <span style="font-size: 12px;">{impacto_logistica}</span>
                        </div>

                        <table style="width: 100%; font-size: 12px;">
                            <tr><td>📅 Data:</td><td style="text-align: right;"><strong>{data_acidente}</strong></td></tr>
                            <tr><td>🏴 UF:</td><td style="text-align: right;"><strong>{acidente['uf']}</strong></td></tr>
                            <tr><td>{periodo}:</td><td style="text-align: right;"><strong>{horario}</strong></td></tr>
                            <tr><td>💀 Mortos:</td><td style="text-align: right;"><strong>{mortos}</strong></td></tr>
                            <tr><td>🏥 Feridos Graves:</td><td style="text-align: right;"><strong>{feridos_graves}</strong></td></tr>
                            <tr><td>🩹 Feridos Leves:</td><td style="text-align: right;"><strong>{int(acidente['feridos_leves']) if pd.notna(acidente['feridos_leves']) else 0}</strong></td></tr>
                        </table>

                        <div style="margin-top: 10px; padding: 8px; background: #e3f2fd; border-radius: 4px;">
                            <strong>🚚 CONTEXTO LOGÍSTICO:</strong><br>
                            <span style="font-size: 11px;">{contexto_logistica_acidente}</span>
                        </div>
                    </div>
                    """
                    
                    folium.Marker(
                        location=[lat, lon],
                        popup=folium.Popup(popup_html, max_width=450),
                        tooltip=f"🚨 BR {br} - KM {km} | {gravidade} | 📦 {impacto_logistica.split(' - ')[0]}",
                        icon=folium.Icon(color=cor, icon='fa-car-crash', prefix='fa')
                    ).add_to(fg_acidentes_graves)
                
                except Exception as e:
                    continue
            
            fg_acidentes_graves.add_to(m2)


        # 7. Camada: Todas as Rodovias (Marker com análise logística)
        if coluna_br in df_enriched.columns:
            df_enriched[coluna_br] = df_enriched[coluna_br].fillna('Não informada').astype(str)

            rodovias_acidentes = df_enriched.groupby(coluna_br).agg({
                'id': 'count',
                'mortos': 'sum',
                'feridos_graves': 'sum',
                'feridos_leves': 'sum',
                'ilesos': 'sum'
            }).reset_index()

            rodovias_acidentes['taxa_mortalidade'] = (rodovias_acidentes['mortos'] / rodovias_acidentes['id']) * 100
            rodovias_acidentes['taxa_mortalidade'] = rodovias_acidentes['taxa_mortalidade'].fillna(0)

            rodovias_acidentes['categoria_risco'] = rodovias_acidentes['taxa_mortalidade'].apply(
                lambda x: 'MUITO ALTO' if x > 3 else 'ALTO' if x > 2 else 'MODERADO' if x > 1 else 'BAIXO'
            )

            rodovias_acidentes['recomendacao_logistica'] = rodovias_acidentes['categoria_risco'].apply(
                lambda x: '🚨 EVITAR' if x == 'MUITO ALTO' else
                         '⚠️ CUIDADO EXTREMO' if x == 'ALTO' else
                         '📋 ATENÇÃO' if x == 'MODERADO' else
                         '✅ ADEQUADO'
            )
            
            fg_todas_rodovias = folium.FeatureGroup(name='🛣️ Todas as Rodovias (Risco Geral)', show=True)
            
            for idx, rodovia in rodovias_acidentes.iterrows():
                try:
                    br = str(rodovia[coluna_br]).split('.')[0] if '.' in str(rodovia[coluna_br]) else str(rodovia[coluna_br])
                    if br == 'Não informada': continue

                    acidentes = int(rodovia['id'])
                    mortos = int(rodovia['mortos'])
                    taxa_mortalidade = float(rodovia['taxa_mortalidade'])
                    recomendacao = rodovia['recomendacao_logistica']
                    categoria = rodovia['categoria_risco']

                    if taxa_mortalidade > 3:
                        cor_rodovia, risco_rodovia, cor_bg = 'darkred', "MUITO ALTA", "#ffebee"
                    elif taxa_mortalidade > 2:
                        cor_rodovia, risco_rodovia, cor_bg = 'red', "ALTA", "#fff3e0"
                    elif taxa_mortalidade > 1:
                        cor_rodovia, risco_rodovia, cor_bg = 'orange', "MÉDIA", "#fff8e1"
                    # Lógica para rotas seguras (Zero Mortalidade)
                    elif taxa_mortalidade == 0:
                        # Adicionar marcador positivo na camada de rotas seguras (fg_rotas_seguras)
                        acidentes_br = coords_validas[coords_validas[coluna_br] == br]
                        if len(acidentes_br) > 0:
                            lat_media = float(acidentes_br['lat_clean'].mean())
                            lon_media = float(acidentes_br['lon_clean'].mean())
                            
                            popup_seguro = f"""
                            <div style="font-family: Arial; min-width: 250px; background: white; color: black; padding: 15px; border-radius: 8px; border: 2px solid green;">
                                <h4 style="margin: 0 0 10px 0; color: #2ecc71;">✅ BR {br} - ROTA SEGURA (RODOVIA)</h4>
                                <div style="background: #2ecc71; color: white; padding: 5px; border-radius: 4px; text-align: center; margin-bottom: 10px;">
                                    <strong>Taxa de Mortalidade: 0.00%</strong>
                                </div>
                                <table style="width: 100%; font-size: 12px;">
                                    <tr><td>📊 Acidentes:</td><td style="text-align: right;"><strong>{acidentes:,}</strong></td></tr>
                                    <tr><td>📦 Recomendação Logística:</td><td style="text-align: right; color: #2ecc71;"><strong>RISCO BAIXO / ADEQUADO</strong></td></tr>
                                </table>
                            </div>
                            """
                            folium.Marker(
                                location=[lat_media, lon_media],
                                popup=folium.Popup(popup_seguro, max_width=300),
                                tooltip=f"✅ BR {br}: Rota Segura (0 mortes) | {acidentes:,} acidentes",
                                icon=folium.Icon(color='green', icon='fa-thumbs-up', prefix='fa'),
                            ).add_to(fg_rotas_seguras)

                        # Usar a cor de risco baixo para a camada de Rodovias Geral (fg_todas_rodovias)
                        cor_rodovia, risco_rodovia, cor_bg = 'green', "BAIXA", "#e8f5e8"
                    
                    else: # Taxa entre 0 e 1% (Baixo Risco)
                        cor_rodovia, risco_rodovia, cor_bg = 'green', "BAIXA", "#e8f5e8"


                    acidentes_br = coords_validas[coords_validas[coluna_br] == br]

                    if len(acidentes_br) > 0:
                        lat_media = float(acidentes_br['lat_clean'].mean())
                        lon_media = float(acidentes_br['lon_clean'].mean())
                        
                        logistica_impacto = (
                            "Alto impacto nas operações - buscar alternativas" if taxa_mortalidade > 2 else
                            "Impacto moderado - avaliar custo-benefício" if taxa_mortalidade > 1 else
                            "Baixo impacto - condições aceitáveis para logística"
                        )

                        popup_content = f"""
                        <div style="font-family: Arial; min-width: 320px; background: {cor_bg}; padding: 15px; border-radius: 8px; border: 2px solid {cor_rodovia};">
                            <h4 style="margin: 0 0 10px 0; color: {cor_rodovia};">🛣️ BR {br}</h4>
                            <div style="background: {cor_rodovia}; color: white; padding: 8px; border-radius: 4px; text-align: center; margin-bottom: 10px;">
                                <strong>PERICULOSIDADE: {risco_rodovia}</strong>
                            </div>

                            <div style="background: #2196f3; color: white; padding: 6px; border-radius: 4px; text-align: center; margin-bottom: 10px;">
                                <strong>📦 {recomendacao}</strong>
                            </div>

                            <table style="width: 100%; font-size: 12px;">
                                <tr><td>📊 Total de Acidentes:</td><td style="text-align: right;"><strong>{acidentes:,}</strong></td></tr>
                                <tr><td>💀 Mortes:</td><td style="text-align: right;"><strong>{mortos:,}</strong></td></tr>
                                <tr><td>📈 Taxa de Mortalidade:</td><td style="text-align: right;"><strong>{taxa_mortalidade:.2f}%</strong></td></tr>
                            </table>

                            <div style="margin-top: 10px; padding: 8px; background: #e3f2fd; border-radius: 4px;">
                                <strong>🚚 IMPACTO LOGÍSTICO GERAL:</strong><br>
                                <span style="font-size: 11px;">{logistica_impacto}</span>
                            </div>
                        </div>
                        """
                        # Adicionar o Marker apenas se não for uma rodovia 'Não informada'
                        folium.Marker(
                            location=[lat_media, lon_media],
                            popup=folium.Popup(popup_content, max_width=400),
                            tooltip=f"🛣️ BR {br}: {acidentes:,} acidentes | {taxa_mortalidade:.2f}% mortalidade | 📦 {recomendacao}",
                            icon=folium.Icon(color=cor_rodovia, icon='fa-road', prefix='fa')
                        ).add_to(fg_todas_rodovias)

                except Exception as e:
                    continue

            fg_todas_rodovias.add_to(m2)


        # 8. Adicionar controles interativos 
        plugins.Fullscreen(position="topright").add_to(m2)
        plugins.MiniMap(tile_layer="CartoDB positron", position="bottomright").add_to(m2)
        # <<< NOVO: ADIÇÃO DOS PLUGINS DO MAPA BÁSICO >>>
        plugins.LocateControl(position="topright").add_to(m2)
        plugins.MeasureControl(position="topleft").add_to(m2)
        
        # O LayerControl irá mostrar TODAS as camadas
        folium.LayerControl(collapsed=False).add_to(m2)

        # 9. Título e Legenda HTML (permanecem os mesmos)
        title_html2 = f'''
        <div style="
            position: fixed;
            top: 10px;
            left: 50%;
            transform: translateX(-50%);
            z-index: 9999;
            background: rgba(44, 62, 80, 0.9);
            padding: 12px 25px;
            border: 2px solid #e74c3c;
            border-radius: 8px;
            box-shadow: 0 4px 15px rgba(0,0,0,0.3);
            text-align: center;
            backdrop-filter: blur(5px);
        ">
            <h4 style="margin: 0; color: white; font-size: 16px;">💀 MAPA DE MORTALIDADE + 📦 ANÁLISE DE LOGÍSTICA</h4>
            <p style="margin: 5px 0 0 0; font-size: 12px; color: #ecf0f1;">Análise logística em todos os pontos</p>
        </div>
        '''
        m2.get_root().html.add_child(folium.Element(title_html2))

        legend_html2 = '''
        <div style="
            position: fixed;
            bottom: 50px;
            left: 50px;
            width: 400px;
            height: auto;
            background: rgba(44, 62, 80, 0.95);
            border: 2px solid #34495e;
            z-index: 9999;
            font-size: 12px;
            padding: 15px;
            border-radius: 8px;
            box-shadow: 0 4px 15px rgba(0,0,0,0.3);
            color: white;
            backdrop-filter: blur(5px);
        ">
            <h4 style="margin: 0 0 12px 0; text-align: center; color: #e74c3c;">📊 LEGENDA + 📦 LOGÍSTICA</h4>

            <div style="margin-bottom: 10px;">
                <strong>🎯 Níveis de Mortalidade (Estados/Rodovias):</strong>
                <div style="display: flex; align-items: center; margin: 5px 0;">
                    <div style="width: 12px; height: 12px; background: darkred; border-radius: 50%; margin-right: 8px;"></div>
                    <span>MUITO ALTO (>3%) - 🚨 EVITAR</span>
                </div>
                <div style="display: flex; align-items: center; margin: 5px 0;">
                    <div style="width: 12px; height: 12px; background: red; border-radius: 50%; margin-right: 8px;"></div>
                    <span>ALTO (2-3%) - ⚠️ CUIDADO EXTREMO</span>
                </div>
                <div style="display: flex; align-items: center; margin: 5px 0;">
                    <div style="width: 12px; height: 12px; background: orange; border-radius: 50%; margin-right: 8px;"></div>
                    <span>MÉDIO (1-2%) - 📋 ATENÇÃO</span>
                </div>
                <div style="display: flex; align-items: center; margin: 5px 0;">
                    <div style="width: 12px; height: 12px; background: green; border-radius: 50%; margin-right: 8px;"></div>
                    <span>BAIXO (<1%) - ✅ ADEQUADO</span>
                </div>
                
                <hr style="border-color: #7f8c8d; margin: 10px 0;">
                
                <div style="display: flex; align-items: center; margin: 5px 0;">
                    <div style="width: 15px; height: 15px; background: green; border-radius: 50%; margin-right: 8px; border: 2px solid white;"></div>
                    <span><strong>✅ ROTA SEGURA (Mortalidade Zero)</strong></span>
                </div>

            </div>

            <div style="margin-bottom: 10px;">
                <strong>📦 IMPACTO LOGÍSTICO (Acidentes Pontuais):</strong>
                <div style="display: flex; align-items: center; margin: 5px 0;">
                    <div style="width: 15px; height: 15px; background: black; border-radius: 50%; margin-right: 8px;"></div>
                    <span>🔴 PARALISAÇÃO TOTAL (Fatal)</span>
                </div>
                <div style="display: flex; align-items: center; margin: 5px 0;">
                    <div style="width: 15px; height: 15px; background: red; border-radius: 50%; margin-right: 8px;"></div>
                    <span>🟡 ALTO IMPACTO (Feridos Graves)</span>
                </div>
            </div>

            <hr style="border-color: #7f8c8d; margin: 10px 0;">

            <div style="font-size: 11px; color: #bdc3c7;">
                <p>🎮 <strong>Como usar para Logística:</strong></p>
                <p>• Clique em QUALQUER ponto para ver análise logística detalhada</p>
                <p>• Use o controle de camadas para alternar <strong>Rotas Seguras</strong>, Rodovias e Acidentes Pontuais</p>
            </div>
        </div>
        '''
        m2.get_root().html.add_child(folium.Element(legend_html2))
        
        return m2

    # ==================== TABELAS - SEM CACHE ====================
    def create_metrics_table(_self):
        """Tabela resumo com as principais métricas do dataset"""
        df_copy = _self.df.copy()
        total_acidentes = len(df_copy)
        total_mortos = df_copy['mortos'].sum()
        total_feridos_graves = df_copy['feridos_graves'].sum()
        
        metrics = {
            "Total de Acidentes": f"{total_acidentes:,}",
            "Total de Mortos": f"{total_mortos:,}",
            "Total de Feridos Graves": f"{total_feridos_graves:,}",
            "Período Analisado": f"{df_copy['ano'].min()} - {df_copy['ano'].max()}",
            "Estados com Dados": f"{df_copy['uf'].nunique()}",
            "Taxa Média de Mortalidade": f"{(total_mortos / total_acidentes * 100):.2f}%" if total_acidentes > 0 else "0.00%",
            "Média de Mortos por Acidente": f"{(total_mortos / total_acidentes):.3f}" if total_acidentes > 0 else "0.000",
            "Rodovias Analisadas": f"{df_copy['br'].nunique()}",
        }
        return pd.DataFrame(list(metrics.items()), columns=["Métrica", "Valor"])

    def create_highways_table(_self):
        """Ranking das rodovias mais perigosas"""
        if _self.df.empty:
            return pd.DataFrame(columns=["Rodovia", "Acidentes", "Mortos", "Feridos Graves", "Taxa Mortalidade (%)"])
            
        rodovias = _self.df["br"].value_counts().head(10).index.tolist()
        dados = []

        for br in rodovias:
            dados_br = _self.df[_self.df["br"] == br]
            total_acidentes = len(dados_br)
            total_mortos = dados_br["mortos"].sum()
            
            taxa_mortalidade = (
                (total_mortos / total_acidentes * 100) if total_acidentes > 0 else 0
            )
            
            dados.append(
                {
                    "Rodovia": br,
                    "Acidentes": total_acidentes,
                    "Mortos": total_mortos,
                    "Feridos Graves": dados_br["feridos_graves"].sum(),
                    "Taxa Mortalidade (%)": taxa_mortalidade,
                }
            )

        df_rodovias = pd.DataFrame(dados)
        df_rodovias["Taxa Mortalidade (%)"] = df_rodovias["Taxa Mortalidade (%)"].round(
            2
        )
        return df_rodovias.sort_values("Acidentes", ascending=False)

# ==============================================================================
# FUNÇÃO PRINCIPAL STREAMLIT - VERSÃO CORRIGIDA E MELHORADA
# ==============================================================================
def main():
    st.set_page_config(
        page_title="Análise de Acidentes Rodoviários", 
        page_icon="🚗", 
        layout="wide",
        initial_sidebar_state="expanded"
    )

    # CSS personalizado
    st.markdown("""
        <style>
        .main-header {
            color: #1f77b4; 
            text-align: center;
            font-size: 2.5em;
            margin-bottom: 20px;
        }
        .stTabs [data-baseweb="tab-list"] {
            gap: 15px;
        }
        .stTabs [data-baseweb="tab"] {
            height: 50px;
            white-space: nowrap;
            background-color: #f0f2f6;
            border-radius: 4px 4px 0 0;
            padding: 10px 20px;
            font-size: 16px;
            font-weight: bold;
            color: #4a4a4a;
        }
        .stTabs [aria-selected="true"] {
            background-color: #1f77b4;
            color: white;
            border-bottom: 3px solid #1f77b4;
        }
        </style>
    """, unsafe_allow_html=True)
    st.markdown('<h1 class="main-header">🚗 Análise de Acidentes Rodoviários</h1>', unsafe_allow_html=True)

    # Carregar dados
    df, estados_coords = load_data()
    
    # ==============================================================================
    # BARRA LATERAL COM FILTROS ROBUSTOS
    # ==============================================================================

    st.sidebar.title("⚙️ Configurações e Filtros")
    
    df_filtrado_base = df.copy() 
    
    anos_disponiveis = sorted(df_filtrado_base['ano'].unique())
    meses_disponiveis = sorted(df_filtrado_base['mes'].unique())
    tipos_acidente = df_filtrado_base['tipo_acidente'].unique().tolist()
    
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 🔍 Refinamento de Data e Hora")
    
    st.sidebar.markdown("**Seleção de Anos**")
    todos_anos = st.sidebar.checkbox("Selecionar Todos os Anos", value=True, key='chk_anos')
    anos_selecionados = anos_disponiveis if todos_anos else st.sidebar.multiselect("Selecione os anos:", options=anos_disponiveis, default=anos_disponiveis, key='multi_anos')
    
    st.sidebar.markdown("**Seleção de Meses**")
    todos_meses = st.sidebar.checkbox("Selecionar Todos os Meses", value=True, key='chk_meses')
    meses_selecionados = meses_disponiveis if todos_meses else st.sidebar.multiselect("Selecione os meses:", options=meses_disponiveis, default=meses_disponiveis, format_func=lambda x: f"{x:02d}", key='multi_meses')
    
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 🏛️ Refinamento Geográfico")
    todos_estados = st.sidebar.checkbox("Selecionar Todos os Estados", value=True, key='chk_estados')
    estados_selecionados = list(estados_coords.keys()) if todos_estados else st.sidebar.multiselect("Selecione os estados:", options=list(estados_coords.keys()), default=list(estados_coords.keys()), key='multi_estados')
    
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 🚨 Refinamento de Tipos de Acidente")
    todos_tipos = st.sidebar.checkbox("Selecionar Todos os Tipos", value=True, key='chk_tipos')
    tipos_selecionados = tipos_acidente if todos_tipos else st.sidebar.multiselect("Selecione os tipos de acidente:", options=tipos_acidente, default=tipos_acidente, key='multi_tipos')

    df_filtrado = df_filtrado_base[
        (df_filtrado_base['ano'].isin(anos_selecionados)) & 
        (df_filtrado_base['mes'].isin(meses_selecionados)) & 
        (df_filtrado_base['uf'].isin(estados_selecionados)) &
        (df_filtrado_base['tipo_acidente'].isin(tipos_selecionados))
    ].copy()
    
    st.sidebar.markdown(f"**📊 Dados Filtrados:** {len(df_filtrado):,} registros")
    
    if df_filtrado.empty:
        st.warning("⚠️ O refinamento dos filtros resultou em um DataFrame vazio. Ajuste as seleções para continuar.")
        return
        
    st.sidebar.markdown("---")
    autor = st.sidebar.text_input("Autor:", "Equipe de Análise")
    
    altura_mapa_completo = st.sidebar.slider("Altura do Mapa Completo", 500, 1000, 600)
    amostra_mapa = st.sidebar.slider("Amostra para Heatmap", 500, 2000, 1000)
    
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 📄 Conteúdo do Relatório DOCX")
    
    include_evolution = st.sidebar.checkbox("Incluir Evolução Temporal", value=True, key='rep_evolution')
    include_states = st.sidebar.checkbox("Incluir Análise por Estado", value=True, key='rep_states')
    include_types = st.sidebar.checkbox("Incluir Distribuição por Tipo", value=True, key='rep_types')
    include_weekday = st.sidebar.checkbox("Incluir Padrão Semanal", value=True, key='rep_weekday')
    include_highways = st.sidebar.checkbox("Incluir Ranking de Rodovias", value=True, key='rep_highways')
    include_metrics = st.sidebar.checkbox("Incluir Tabela de Métricas", value=True, key='rep_metrics')
    # Removido include_map pois não há mais mapa básico
    include_complete_map = st.sidebar.checkbox("Incluir Ref. ao Mapa Completo", value=True, key='rep_complete_map')
    
    selecoes = {
        "include_evolution": include_evolution,
        "include_states": include_states,
        "include_types": include_types,
        "include_weekday": include_weekday,
        "include_highways": include_highways,
        "include_metrics": include_metrics,
        "include_map": False, # Desativado permanentemente
        "include_complete_map": include_complete_map,
    }

    analyzer = DataAnalyzer(df_filtrado, estados_coords)

    # <<< ALTERAÇÃO 1: Removida a aba "Mapas Básicos" e a variável tab2 >>>
    tab1, tab2, tab3, tab4 = st.tabs(
        ["📊 Análises Gráficas", "🗺️ Mapa", "📈 Métricas e Tabelas", "📋 Resumo Executivo"]
    )

    with tab1:
        st.markdown("## 📊 Análises Gráficas dos Acidentes Rodoviários")
        
        col1, col2 = st.columns(2)
        with col1:
            if selecoes["include_evolution"]:
                st.markdown("### 📈 Evolução Temporal dos Acidentes")
                fig = analyzer.create_evolution_chart()
                st.pyplot(fig)

            if selecoes["include_states"]:
                st.markdown("### 🏛️ Comparativo entre Estados")
                fig = analyzer.create_states_chart()
                st.pyplot(fig)
        
        with col2:
            if selecoes["include_types"]:
                st.markdown("### 🚨 Distribuição por Tipo de Acidente")
                fig = analyzer.create_accident_types_chart()
                st.pyplot(fig)

            if selecoes["include_weekday"]:
                st.markdown("### 📅 Padrão Semanal de Acidentes")
                fig = analyzer.create_weekday_chart()
                st.pyplot(fig)

    # <<< ALTERAÇÃO 2: Bloco "with tab2:" (mapa básico) foi COMPLETAMENTE REMOVIDO >>>

    # <<< ALTERAÇÃO 3: As abas seguintes foram renomeadas (tab3->tab2, tab4->tab3, etc.) >>>
    with tab2: # Antiga tab3
        st.markdown("## 🗺️ Mapa Unificado com Análise de Logística")

        if selecoes["include_complete_map"]:
            st.markdown("### 📦 Mapa Completo de Logística e Mortalidade (Avançado)")
            
            with st.spinner("Gerando mapa unificado..."):
                mapa_completo = analyzer.create_complete_logistics_map(
                    df_filtrado,
                    sample_size=amostra_mapa,
                    map_height=altura_mapa_completo
                )
                folium_static(mapa_completo, width=1000, height=altura_mapa_completo)

    with tab3: # Antiga tab4
        st.markdown("## 📈 Métricas e Análises Consolidadas")

        if selecoes["include_metrics"]:
            st.markdown("### 📊 Métricas Gerais do Período")
            tabela = analyzer.create_metrics_table()
            st.dataframe(tabela, use_container_width=True)

        if selecoes["include_highways"]:
            st.markdown("### 🛣️ Ranking das Rodovias Mais Perigosas")
            tabela = analyzer.create_highways_table()
            st.dataframe(tabela, use_container_width=True)

    with tab4: # Antiga tab5
        st.markdown("## 📋 Resumo Executivo")
        
        resumo_executivo = st.text_area(
            "Resumo Executivo (Introdução)", 
            "A análise abrange o período de {} a {}, identificando que o principal desafio de segurança viária está concentrado nas BRs 101, 116 e 040, com ênfase no tipo 'Colisão Traseira' como o mais frequente.".format(df_filtrado['ano'].min(), df_filtrado['ano'].max())
        )
        
        insights = st.text_area(
            "Insights e Recomendações Estratégicas", 
            "Recomenda-se a implementação de campanhas de conscientização focadas em direção defensiva e distância de segurança, especialmente nas sextas-feiras e sábados (dias de pico de acidentes)."
        )

    # Geração de Relatórios (sem alterações aqui)
    st.markdown("---")
    st.markdown("## 📄 Gerar Relatório Executivo Word (DOCX)")

    if st.button("📥 Gerar Relatório Executivo DOCX", type="primary"):
        with st.spinner("Gerando relatório executivo Word (.docx)..."):
            figuras_graficos = {}
            figuras_mapas = {}
            
            try:
                if selecoes["include_evolution"]: figuras_graficos["evolution"] = analyzer.create_evolution_chart()
                if selecoes["include_states"]: figuras_graficos["states"] = analyzer.create_states_chart()
                if selecoes["include_types"]: figuras_graficos["types"] = analyzer.create_accident_types_chart()
                if selecoes["include_weekday"]: figuras_graficos["weekday"] = analyzer.create_weekday_chart()
                
                report = DOCXReportGenerator()
                report.build_report(
                    analyzer, 
                    selecoes, 
                    autor, 
                    resumo_executivo, 
                    insights,
                    figuras_graficos,
                    figuras_mapas
                )
                docx_buffer = report.generate_docx()

                st.download_button(
                    label="📥 Download do Relatório DOCX",
                    data=docx_buffer,
                    file_name=f"relatorio_executivo_acidentes_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                st.success("Relatório executivo Word (.docx) gerado com sucesso!")
                
            except Exception as e:
                st.error(f"Erro ao gerar relatório: {str(e)}")


    # Rodapé
    st.markdown("---")
    st.markdown(
        """
        <div style="text-align: center; color: #666; font-size: 12px;">
        <p>🚗 Análise de Acidentes Rodoviários - Desenvolvido para auxiliar na tomada de decisões estratégicas</p>
        <p>⚠️ Dados simulados para fins de demonstração | 📅 Período: {}-{} | 📊 Total de registros: {:,}</p>
        </div>
        """.format(
            df_filtrado['ano'].min(), 
            df_filtrado['ano'].max(), 
            len(df_filtrado)
        ),
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()